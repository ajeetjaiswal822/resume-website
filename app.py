import os
import pdfplumber
import docx2txt
import json
from flask import Flask, render_template, request, send_file
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

# -------------------------
# LOAD ENV VARIABLES
# -------------------------

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    raise ValueError("OPENAI_API_KEY not found in .env file")

client = OpenAI(api_key=api_key)

# -------------------------
# FLASK SETUP
# -------------------------

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# -------------------------
# READ PDF
# -------------------------

def read_pdf(path):

    text = ""

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    return text


# -------------------------
# READ DOCX
# -------------------------

def read_docx(path):

    return docx2txt.process(path)


# -------------------------
# PARSE RESUME USING AI
# -------------------------

def parse_resume(resume_text):

    prompt = f"""
Extract resume information.

Remove phone numbers and emails.

Return JSON in this format:

{{
"name":"",
"location":"",
"summary":[
 "point1",
 "point2"
],
"education":[
 {{
   "degree":"",
   "field":"",
   "institution":"",
   "duration":""
 }}
],
"certifications":[
 {{
   "name":"",
   "year":""
 }}
],
"skills":[],
"experience":[
 {{
   "company":"",
   "location":"",
   "title":"",
   "duration":"",
   "responsibilities":[]
 }}
]
}}

Rules:
- Professional summary must be converted into 4-6 bullet points
- Keep chronological order
- Do not merge jobs
- Responsibilities must be bullet points

Resume:
{resume_text[:12000]}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
        temperature=0
    )

    return json.loads(response.choices[0].message.content)


# -------------------------
# PRESCREEN ANSWERS
# -------------------------

def answer_prescreen(resume_text, questions):

    questions_text = "\n".join(questions)

    prompt = f"""
Answer the prescreen questions based ONLY on the resume.

If answer is not found say:
"Not mentioned in resume"

Resume:
{resume_text[:12000]}

Questions:
{questions_text}

Return JSON:

{{
"answers":[
 {{
   "question":"",
   "answer":""
 }}
]
}}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
        temperature=0
    )

    return json.loads(response.choices[0].message.content)


# -------------------------
# FORMAT BODY TEXT
# -------------------------

def format_body(paragraph):

    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


# -------------------------
# FORMAT HEADINGS
# -------------------------

def format_heading(paragraph):

    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.bold = True

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1
    paragraph_format.space_before = Pt(10)
    paragraph_format.space_after = Pt(2)


# -------------------------
# CREATE DOCX
# -------------------------

def create_docx(data, filename, prescreen=None):

    doc = Document()

    # Narrow margins
    section = doc.sections[0]

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Name
    heading = doc.add_heading(data.get("name",""), level=0)
    heading.alignment = 1
    format_body(heading)

    # Location
    location = doc.add_paragraph(data.get("location",""))
    location.alignment = 1
    format_body(location)

    # Prescreen
    if prescreen:

        h = doc.add_heading("Prescreen", level=1)
        format_heading(h)

        for qa in prescreen:

            q = doc.add_paragraph("Q: " + qa["question"])
            format_body(q)

            a = doc.add_paragraph("A: " + qa["answer"])
            format_body(a)

    # Professional Summary
    h = doc.add_heading("Professional Summary", level=1)
    format_heading(h)

    for point in data.get("summary", []):
        p = doc.add_paragraph(point, style="List Bullet")
        format_body(p)

    # Education
    h = doc.add_heading("Education", level=1)
    format_heading(h)

    for edu in data.get("education", []):

        degree = edu.get("degree","")
        field = edu.get("field","")
        institution = edu.get("institution","")
        duration = edu.get("duration","")

        line = f"{degree} in {field} – {institution} ({duration})"

        p = doc.add_paragraph(line)
        format_body(p)

    # Certifications
    h = doc.add_heading("Certification", level=1)
    format_heading(h)

    for cert in data.get("certifications", []):

        name = cert.get("name","")
        year = cert.get("year","")

        line = f"{name} ({year})"

        p = doc.add_paragraph(line)
        format_body(p)

    # Skills
    h = doc.add_heading("Skills", level=1)
    format_heading(h)

    for skill in data.get("skills", []):

        p = doc.add_paragraph(skill, style="List Bullet")
        format_body(p)

    # Experience
    h = doc.add_heading("Professional Experience", level=1)
    format_heading(h)

    for job in data.get("experience", []):

        company = job.get("company","")
        location = job.get("location","")
        title = job.get("title","")
        duration = job.get("duration","")

        company_line = f"{company} – {location}    {duration}"

        p = doc.add_paragraph()
        run = p.add_run(company_line)
        run.bold = True
        format_body(p)

        p2 = doc.add_paragraph(title)
        format_body(p2)

        for resp in job.get("responsibilities", []):
            p = doc.add_paragraph(resp, style="List Bullet")
            format_body(p)

        doc.add_paragraph("")

    output_file = os.path.join(OUTPUT_FOLDER, filename + "_formatted.docx")

    doc.save(output_file)

    return output_file


# -------------------------
# WEB ROUTES
# -------------------------

@app.route("/")
def home():

    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():

    file = request.files["resume"]
    prescreen_text = request.form.get("prescreen","")

    filename = file.filename
    input_path = os.path.join(UPLOAD_FOLDER, filename)

    file.save(input_path)

    try:

        if filename.endswith(".pdf"):
            text = read_pdf(input_path)

        elif filename.endswith(".docx") or filename.endswith(".doc"):
            text = read_docx(input_path)

        else:
            return "Unsupported file format"

        parsed = parse_resume(text)

        prescreen_answers = None

        if prescreen_text.strip():

            questions = [q.strip() for q in prescreen_text.split("\n") if q.strip()]

            prescreen_data = answer_prescreen(text, questions)

            prescreen_answers = prescreen_data["answers"]

        output_file = create_docx(parsed, filename, prescreen_answers)

        return send_file(output_file, as_attachment=True)

    except Exception as e:

        return str(e)


# -------------------------
# RUN SERVER
# -------------------------

if __name__ == "__main__":

    app.run()